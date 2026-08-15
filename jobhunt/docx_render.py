"""The .docx resume renderer (Phase 2). A validated `select.Selection` -> a .docx.

Isolated from the select/validate engine and from the RenderCV path: a new output
format is a new claim surface, so it inherits the whole honesty backbone. It emits
ONLY verbatim library / validated-Selection strings and structured fields
(identity from facts, education from the library) — no interpolation, no
reformatted number, no re-derived descriptor. If a value is not from the library
or the Selection, it does not appear.

This is the §8 home for the work-authorization header line that RenderCV had no
slot for: python-docx gives real header control, so the line is placed as its own
contact line rather than stuffed into the summary.

python-docx is imported lazily (the `[resume]` extra), so importing this module
does not require it until a docx is actually built.
"""

from __future__ import annotations

import io
import sqlite3
from typing import Any

from . import queries

# entry_key -> the repos_public key that gates that project's link (strict is True).
_REPO_KEY = {"ARC": "ARC", "JOB": "jobhunt", "PEP": "peptideDesign", "FPS": "FireProofSheep"}

# §8 headings, literal and in order.
_SECTION_ORDER = ("summary", "skills", "experience", "projects", "education")
_HEADINGS = {
    "summary": "Summary",
    "skills": "Technical Skills",
    "experience": "Experience",
    "projects": "Projects",
    "education": "Education",
}


class DocxError(RuntimeError):
    """The selection could not be rendered to a .docx."""


def _contact_line(facts: dict[str, str]) -> str:
    """Location · email · phone · website · linkedin · github — in the body, §8."""
    parts: list[str] = []
    city, state = facts.get("identity.city"), facts.get("identity.state")
    if city or state:
        parts.append(", ".join(p for p in (city, state) if p))
    for key in ("identity.email", "identity.phone", "identity.website"):
        if facts.get(key):
            parts.append(facts[key])
    for key in ("identity.linkedin", "identity.github"):
        url = (facts.get(key) or "").rstrip("/")
        if url:
            parts.append(url.split("//", 1)[-1])  # display without the scheme
    return "  ·  ".join(parts)


def build_docx(conn: sqlite3.Connection, selection: Any) -> bytes:
    """Render a validated `select.Selection` to .docx bytes, per §8."""
    try:
        from docx import Document
        from docx.enum.text import WD_TAB_ALIGNMENT
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover - depends on the [resume] extra
        raise DocxError(
            "python-docx is not installed. Install it with: uv pip install -e '.[resume]'"
        ) from exc

    facts = queries.profile_facts(conn)
    name = facts.get("identity.preferred_name") or facts.get("identity.legal_name")
    if not name:
        raise DocxError("no name in profile_facts — fill identity in facts.yaml")

    doc = Document()
    section = doc.sections[0]
    for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, side, Inches(0.6))
    # A right tab stop at the text width, so dates align right without spaces (§8).
    right_edge = section.page_width - section.left_margin - section.right_margin
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(10.5)

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        # A rule under the heading, in the theme's spirit — a bottom border.
        pPr = p._p.get_or_add_pPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "888888")):
            bottom.set(qn(k), v)
        borders.append(bottom)
        pPr.append(borders)

    def dated_line(left_bold: str, left_rest: str, date_text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(right_edge, WD_TAB_ALIGNMENT.RIGHT)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(left_bold)
        run.bold = True
        if left_rest:
            p.add_run(left_rest)
        if date_text:
            p.add_run("\t" + date_text)

    def bullet(text: str) -> None:
        # A real list bullet via list formatting — never a literal • in a run (§8).
        doc.add_paragraph(text, style="List Bullet")

    # ---- header: name, contact line (body), work-auth line (§8) ----
    hp = doc.add_paragraph()
    hrun = hp.add_run(name)
    hrun.bold = True
    hrun.font.size = Pt(20)
    hp.paragraph_format.space_after = Pt(2)

    contact = _contact_line(facts)
    if contact:
        cp = doc.add_paragraph(contact)
        cp.paragraph_format.space_after = Pt(0)
    work_auth = facts.get("identity.work_authorization_line")
    if work_auth:
        wp = doc.add_paragraph()
        wr = wp.add_run(work_auth)
        wr.font.size = Pt(10)
        wr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        wp.paragraph_format.space_after = Pt(2)

    # ---- structured content, placed by entry_key (same rules as the RenderCV path) ----
    by_entry: dict[str, list[Any]] = {}
    order: list[str] = []
    for b in selection.bullets:
        by_entry.setdefault(b.entry_key, []).append(b)
        if b.entry_key not in order:
            order.append(b.entry_key)

    variant_entries = queries.resume_variant_entries(conn, selection.variant)
    swap = queries.resume_swap_entries(conn)
    by_key = {ve["entry_key"]: ve for ve in variant_entries}
    repos_public = (queries.resume_meta(conn, "open_facts") or {}).get("repos_public") or {}

    for kind in _SECTION_ORDER:
        if kind == "summary":
            if selection.summary and selection.summary.strip():
                heading(_HEADINGS["summary"])
                doc.add_paragraph(selection.summary.strip())
        elif kind == "skills":
            rows = [(g, ", ".join(items)) for g, items in selection.skills if items]
            if rows:
                heading(_HEADINGS["skills"])
                for group, items in rows:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.add_run(f"{group}: ").bold = True
                    p.add_run(items)
        elif kind == "education":
            edu = queries.resume_education(conn)
            if edu:
                heading(_HEADINGS["education"])
                for e in edu:
                    dated_line(str(e["school"]), f", {e['degree']}", str(e["date_text"]))
        elif kind == "experience":
            entries = [ve for ve in variant_entries
                       if ve["kind"] == "experience" and by_entry.get(ve["entry_key"])]
            if entries:
                heading(_HEADINGS["experience"])
                for ve in entries:
                    dated_line(str(ve["company"]), f", {ve['role']}", str(ve["date_text"]))
                    for b in by_entry[ve["entry_key"]]:
                        bullet(b.text)
        elif kind == "projects":
            variant_proj = [ve["entry_key"] for ve in variant_entries if ve["kind"] == "project"]
            present = [k for k in order if k in by_key and by_key[k]["kind"] == "project"] + \
                      [k for k in order if k in swap]
            ordered = [k for k in variant_proj if k in present] + \
                      [k for k in present if k not in variant_proj]
            ordered = ordered[:2]
            if ordered:
                heading(_HEADINGS["projects"])
                for entry_key in ordered:
                    chosen = by_entry.get(entry_key)
                    header = by_key.get(entry_key) or swap.get(entry_key)
                    if not chosen or header is None:
                        continue
                    name_text = str(header["name"])
                    # Link only on strict `is True` and only if a URL is recorded.
                    if repos_public.get(_REPO_KEY.get(entry_key)) is True and header.get("url"):
                        name_text = f"{name_text} ({header['url']})"
                    dated_line(name_text, "", str(header["date_text"]))
                    if header.get("descr"):
                        dp = doc.add_paragraph(str(header["descr"]))
                        dp.paragraph_format.space_after = Pt(0)
                        for run in dp.runs:
                            run.italic = True
                    for b in chosen:
                        bullet(b.text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_text_in_order(data: bytes) -> list[str]:
    """Read a .docx back as its paragraph text in document order.

    The copy-paste-in-order check at the master level: a faithful docx -> PDF
    export preserves this order in the text layer, so verifying it on the docx is
    the property §8's copy-paste test is really about.
    """
    from docx import Document

    doc = Document(io.BytesIO(data))
    return [p.text for p in doc.paragraphs if p.text.strip()]


# =============================== docx -> PDF ===============================
#
# The submission PDF is derived from the .docx master via LibreOffice headless
# (the reviewed Phase 2 decision), which carries the §8 work-auth header line
# through to the PDF and produces a text-based, copy-paste-in-order PDF.
# LibreOffice is therefore a deploy dependency on the host; this raises a clear
# error rather than silently falling back to a renderer that drops the work-auth
# line.

_SOFFICE_CANDIDATES = (
    "soffice",
    "libreoffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
)


def find_soffice() -> str | None:
    """The LibreOffice CLI, or None when it is not installed."""
    import os
    import shutil

    for candidate in _SOFFICE_CANDIDATES:
        found = shutil.which(candidate)
        if found:
            return found
        if os.path.isfile(candidate) and os.access(candidate, os.X_OK):
            return candidate
    return None


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert .docx bytes to a text-based PDF via LibreOffice headless.

    Raises `DocxError` naming the missing dependency rather than falling back —
    Option 1 was chosen precisely so the work-auth line reaches the PDF, and a
    silent fallback to RenderCV would drop it.
    """
    import subprocess
    import tempfile
    from pathlib import Path

    soffice = find_soffice()
    if soffice is None:
        raise DocxError(
            "LibreOffice (soffice) is not installed, so the .docx cannot be "
            "converted to the submission PDF. Install it on the host — e.g. "
            "`brew install --cask libreoffice`. It is the reviewed docx->PDF path."
        )
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "resume.docx"
        src.write_bytes(docx_bytes)
        # A per-conversion user profile dir, so concurrent runs do not collide on
        # LibreOffice's default profile lock.
        profile = Path(tmp) / "lo-profile"
        result = subprocess.run(
            [soffice, "--headless", f"-env:UserInstallation=file://{profile}",
             "--convert-to", "pdf", "--outdir", tmp, str(src)],
            capture_output=True, text=True, timeout=120,
        )
        produced = Path(tmp) / "resume.pdf"
        if result.returncode != 0 or not produced.is_file():
            detail = (result.stdout + result.stderr).strip()[-800:]
            raise DocxError(f"soffice docx->pdf failed (exit {result.returncode}):\n{detail}")
        return produced.read_bytes()


def pdf_text_in_order(pdf_bytes: bytes) -> list[str]:
    """The PDF's text as a flat token stream, for the §8 copy-paste-in-order test.

    Uses pypdf if present; otherwise raises so the check is never silently
    skipped. Run on the host where the PDF (and LibreOffice) exist.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise DocxError(
            "pypdf is not installed, so the PDF copy-paste-in-order check cannot "
            "run. Install it with: uv pip install pypdf"
        ) from exc

    reader = PdfReader(io.BytesIO(pdf_bytes))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return [line.strip() for line in text.splitlines() if line.strip()]


def verify_copy_paste_order(docx_bytes: bytes, pdf_bytes: bytes) -> bool:
    """True if the PDF's text preserves the .docx master's paragraph order.

    The §8 copy-paste-in-order test, verified rather than asserted: every docx
    line appears in the PDF text, in the same relative order.
    """
    pdf_text = "\n".join(pdf_text_in_order(pdf_bytes))
    cursor = 0
    for line in docx_text_in_order(docx_bytes):
        head = line[:40]
        found = pdf_text.find(head, cursor)
        if found < 0:
            return False
        cursor = found + len(head)
    return True
